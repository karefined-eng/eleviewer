from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QTextBrowser, QToolButton,
)
from PySide6.QtCore import Signal, QSize, QUrl

# Lazy WebEngine flag — checked once on first instantiation
_WEB_AVAILABLE: bool | None = None

import base64
import io
import os
import zipfile
import xml.etree.ElementTree as ET

try:
    import mammoth
    MAMMOTH_AVAILABLE = True
except ImportError:
    MAMMOTH_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from icons import icon
from theme import (
    editor_stylesheet, compact_toolbar_stylesheet, ICON_SIZE_COMPACT,
    get_brand_accent, get_active_palette,
)


class DocxViewer(QWidget):
    """
    DOCX viewer and editor.
    Renders rich HTML with embedded Base64 images via mammoth (in QTextBrowser)
    or fallback to python-docx text+image extraction.
    """

    textChanged = Signal()

    def __init__(self, file_path=None):
        super().__init__()

        self.file_path = file_path
        self.is_modified = False
        self.docx_content = None
        self._bookmark_callback = None

        layout = QVBoxLayout()
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        toolbar = QHBoxLayout()
        toolbar.setContentsMargins(4, 4, 4, 4)
        toolbar.addStretch()

        icon_sz = ICON_SIZE_COMPACT
        icon_qsize = QSize(icon_sz, icon_sz)
        self.btn_bookmark = QToolButton()
        self.btn_bookmark.setIconSize(icon_qsize)
        self.btn_bookmark.setIcon(icon("book-open", size=icon_sz))
        self.btn_bookmark.setToolTip("Bookmark this position")
        self.btn_bookmark.setStyleSheet(compact_toolbar_stylesheet())
        self.btn_bookmark.setAutoRaise(True)
        self.btn_bookmark.clicked.connect(self._add_bookmark_here)
        toolbar.addWidget(self.btn_bookmark)

        # ── Lazy-load WebEngine (Rule 25: no module-level import) ────────────
        global _WEB_AVAILABLE
        if _WEB_AVAILABLE is None:
            try:
                from web_panel import WEB_AVAILABLE as _WA
                _WEB_AVAILABLE = _WA
            except Exception:
                _WEB_AVAILABLE = False

        if _WEB_AVAILABLE:
            from web_panel import _SecureWebView
            self.editor = _SecureWebView()
            self._use_webengine = True
        else:
            self.editor = QTextBrowser()
            self.editor.setOpenExternalLinks(True)
            self.editor.setStyleSheet(editor_stylesheet())
            self.editor.textChanged.connect(self._on_text_changed)
            self._use_webengine = False

        layout.addLayout(toolbar)
        layout.addWidget(self.editor)
        self.setLayout(layout)

        if file_path:
            self.load_from_path(file_path)

    def set_bookmark_callback(self, callback):
        self._bookmark_callback = callback

    # ── In-document search (Ctrl+F via global FindReplaceWidget) ──────

    def find_text(self, text, match_case=False, whole_word=False, forward=True):
        """Find text; delegates to WebEngine or QTextBrowser API."""
        if self._use_webengine:
            from PySide6.QtWebEngineCore import QWebEnginePage
            flags = QWebEnginePage.FindFlag(0)
            if not forward:
                flags |= QWebEnginePage.FindFlag.FindBackward
            self.editor.findText(text, flags)
            return True
        from PySide6.QtGui import QTextDocument
        flags = QTextDocument.FindFlag(0)
        if match_case:
            flags |= QTextDocument.FindFlag.FindCaseSensitively
        if whole_word:
            flags |= QTextDocument.FindFlag.FindWholeWords
        if not forward:
            flags |= QTextDocument.FindFlag.FindBackward
        return self.editor.find(text, flags)

    def load_from_path(self, file_path):
        """Load DOCX file from disk with progressive fallbacks."""
        self.file_path = file_path
        if DOCX_AVAILABLE:
            try:
                self.docx_content = Document(file_path)
            except Exception as e:
                print(f"[DOCX Viewer] python-docx load failed: {e}")
                self.docx_content = None
        self._display_content()
        self.is_modified = False

    def _display_content(self):
        """Render DOCX to rich HTML with base64 images via mammoth, falling back to python-docx and stdlib zip parsing."""
        if not self.file_path and not self.docx_content:
            return

        self.editor.blockSignals(True)

        if MAMMOTH_AVAILABLE and self.file_path and os.path.exists(self.file_path):
            try:
                with open(self.file_path, "rb") as docx_file:
                    result = mammoth.convert_to_html(
                        docx_file,
                        convert_image=mammoth.images.img_element(
                            self._mammoth_image_converter
                        ),
                    )
                    html_content = result.value
                    styled_html = self._wrap_html(html_content)
                    if self._use_webengine:
                        self.editor.setHtml(styled_html, QUrl())
                    else:
                        self.editor.setHtml(styled_html)
                    self.editor.blockSignals(False)
                    self.is_modified = False
                    return
            except Exception as e:
                print(f"[DOCX Viewer] Mammoth HTML conversion failed: {e}")

        # Fallback 1: python-docx text + image extraction as HTML
        if self.docx_content:
            try:
                html = self._build_html_from_docx(self.docx_content)
                styled_html = self._wrap_html(html)
                if self._use_webengine:
                    self.editor.setHtml(styled_html, QUrl())
                else:
                    self.editor.setHtml(styled_html)
                self.editor.blockSignals(False)
                self.is_modified = False
                return
            except Exception as e:
                print(f"[DOCX Viewer] python-docx extraction failed: {e}")

        # Fallback 2: Ponytail zero-dependency stdlib ZIP parser
        if self.file_path and os.path.exists(self.file_path):
            html = self._fallback_zip_parse(self.file_path)
            styled_html = self._wrap_html(html)
            if self._use_webengine:
                self.editor.setHtml(styled_html, QUrl())
            else:
                self.editor.setHtml(styled_html)

        self.editor.blockSignals(False)
        self.is_modified = False

    def _fallback_zip_parse(self, file_path):
        """Zero-dependency stdlib fallback for DOCX files using zipfile & ElementTree."""
        try:
            with zipfile.ZipFile(file_path, "r") as z:
                if "word/document.xml" not in z.namelist():
                    return "<i>(Empty or invalid DOCX document)</i>"

                doc_xml = z.read("word/document.xml")
                tree = ET.fromstring(doc_xml)

                # Extract media images
                image_map = {}
                if "word/_rels/document.xml.rels" in z.namelist():
                    rels_tree = ET.fromstring(z.read("word/_rels/document.xml.rels"))
                    for rel in rels_tree.iter():
                        if rel.tag.endswith("}Relationship"):
                            target = rel.attrib.get("Target", "")
                            rel_id = rel.attrib.get("Id", "")
                            if "media/" in target:
                                media_filename = target.split("/")[-1]
                                media_path = f"word/media/{media_filename}"
                                if media_path in z.namelist():
                                    img_bytes = z.read(media_path)
                                    b64 = base64.b64encode(img_bytes).decode("utf-8")
                                    ext = media_filename.split(".")[-1].lower()
                                    mime = "image/jpeg" if ext == "jpg" else f"image/{ext}"
                                    image_map[rel_id] = f'<img src="data:{mime};base64,{b64}" style="max-width:100%; border-radius:4px; margin:8px 0; display:block;" />'

                # Extract text paragraphs and embedded blips
                html_parts = []
                for p in tree.iter():
                    if p.tag.endswith("}p"):
                        p_texts = []
                        for elem in p.iter():
                            if elem.tag.endswith("}t") and elem.text:
                                p_texts.append(elem.text)
                            elif elem.tag.endswith("}blip"):
                                embed_id = elem.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                                if embed_id and embed_id in image_map:
                                    p_texts.append(image_map[embed_id])

                        if p_texts:
                            html_parts.append(f"<p>{''.join(p_texts)}</p>")

                return "".join(html_parts) if html_parts else "<i>(No text content found)</i>"
        except Exception as e:
            return f"<i>(Failed to parse DOCX: {e})</i>"

    @staticmethod
    def _mammoth_image_converter(image):
        """Convert mammoth image to Base64 data URI for inline rendering."""
        with image.open() as image_bytes:
            raw = image_bytes.read()
        content_type = image.content_type or "image/png"
        b64 = base64.b64encode(raw).decode("utf-8")
        return {
            "src": f"data:{content_type};base64,{b64}",
        }

    def _build_html_from_docx(self, doc):
        """Build HTML from python-docx Document with embedded Base64 images."""
        html_parts = []

        # Build a map of relationship IDs to image data for quick lookup
        image_map = {}
        try:
            for rel_id, rel in doc.part.rels.items():
                if "image" in rel.reltype:
                    try:
                        img_bytes = rel.target_part.blob
                        content_type = rel.target_part.content_type or "image/png"
                        b64 = base64.b64encode(img_bytes).decode("utf-8")
                        image_map[rel_id] = (content_type, b64)
                    except Exception:
                        pass
        except Exception:
            pass

        # Namespace shortcuts for XML element search
        ns_drawing = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
        ns_blip = "{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
        ns_embed = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"

        for para in doc.paragraphs:
            # Check paragraph style for heading detection
            style_name = (para.style.name or "").lower() if para.style else ""
            tag = "p"
            if "heading 1" in style_name:
                tag = "h1"
            elif "heading 2" in style_name:
                tag = "h2"
            elif "heading 3" in style_name:
                tag = "h3"
            elif "heading 4" in style_name:
                tag = "h4"

            # Collect text and inline images from runs
            run_html = []
            for run in para.runs:
                # Check if this run contains an embedded image (drawing element)
                drawings = run._element.findall(f".//{ns_drawing}")
                if drawings:
                    for drawing in drawings:
                        blips = drawing.findall(f".//{ns_blip}")
                        for blip in blips:
                            embed_id = blip.get(ns_embed)
                            if embed_id and embed_id in image_map:
                                ct, b64 = image_map[embed_id]
                                run_html.append(
                                    f'<img src="data:{ct};base64,{b64}" '
                                    f'style="max-width:100%; border-radius:4px; '
                                    f'margin:8px 0; display:block;" />'
                                )

                # Add text content
                text = run.text
                if text:
                    # Apply basic formatting
                    if run.bold:
                        text = f"<b>{text}</b>"
                    if run.italic:
                        text = f"<i>{text}</i>"
                    if run.underline:
                        text = f"<u>{text}</u>"
                    run_html.append(text)

            content = "".join(run_html)
            if content.strip():
                html_parts.append(f"<{tag}>{content}</{tag}>")

        # Extract tables
        for table in doc.tables:
            table_html = ['<table style="border-collapse:collapse; width:100%; margin:12px 0;">']
            p = get_active_palette()
            for i, row in enumerate(table.rows):
                cell_tag = "th" if i == 0 else "td"
                row_html = [
                    f'<{cell_tag} style="border:1px solid {p["BRAND_BORDER"]}; padding: 6px;">{cell.text}</{cell_tag}>'
                    for cell in row.cells
                ]
                table_html.append(f"<tr>{''.join(row_html)}</tr>")
            table_html.append("</table>")
            html_parts.append("".join(table_html))

        return "\n".join(html_parts)

    def _wrap_html(self, body_html):
        """Wrap raw HTML content in a styled document shell matching our dark theme."""
        p = get_active_palette()
        accent = get_brand_accent()
        return f"""
        <html><head><style>
            body {{
                background: {p['BRAND_BACKGROUND']};
                color: {p['BRAND_PRIMARY']};
                font-family: 'Segoe UI', sans-serif;
                font-size: 15px;
                line-height: 1.7;
                padding: 16px;
                margin: 0;
            }}
            h1, h2, h3, h4 {{ color: {p['BRAND_PRIMARY']}; margin-top: 1.2em; }}
            h1 {{ font-size: 22px; border-bottom: 1px solid {p['BRAND_BORDER']}; padding-bottom: 8px; }}
            h2 {{ font-size: 18px; }}
            h3 {{ font-size: 16px; }}
            a {{ color: {accent}; }}
            table {{ border-collapse: collapse; width: 100%; margin: 12px 0; }}
            th {{ background: {p['BRAND_PANEL']}; font-weight: bold; }}
            td, th {{ border: 1px solid {p['BRAND_BORDER']}; padding: 6px 10px; }}
            img {{ max-width: 100%; border-radius: 4px; margin: 8px 0; }}
            p {{ margin: 0.5em 0; }}
        </style></head><body>{body_html}</body></html>
        """

    def _on_text_changed(self):
        """Mark as modified when user edits."""
        self.is_modified = True
        self.textChanged.emit()

    def _bookmark_payload(self):
        name = os.path.basename(self.file_path) if self.file_path else "document"
        return {
            "page_number": 0,
            "scroll_position_y": float(self.editor.verticalScrollBar().value()),
            "label": f"Position in {name}",
        }

    def _add_bookmark_here(self):
        if self._bookmark_callback:
            self._bookmark_callback(self._bookmark_payload())

    def go_to_bookmark(self, page_number=0, scroll_position_y=0.0):
        self.editor.verticalScrollBar().setValue(int(scroll_position_y))

    def to_docx_bytes(self):
        """
        Convert current editor content back to DOCX bytes.
        Creates a new document with the edited text.
        """
        try:
            new_doc = Document()
            text = self.editor.toPlainText()
            paragraphs = text.split("\n\n")

            for para_text in paragraphs:
                if para_text.strip():
                    new_doc.add_paragraph(para_text)

            byte_stream = io.BytesIO()
            new_doc.save(byte_stream)
            byte_stream.seek(0)

            return byte_stream.getvalue()

        except Exception as e:
            raise Exception(f"Failed to save DOCX: {str(e)}")

    def toPlainText(self):
        """Compatibility method - returns text content."""
        return self.editor.toPlainText()

    def read_current_page(self, voice_id=None):
        """Duck-typed method for TTS engine. Reads selection or full text."""
        cursor = self.editor.textCursor()
        text = cursor.selectedText()
        if not text:
            text = self.editor.toPlainText()
        return text

    def setPlainText(self, text):
        """Compatibility method - set text content (for reopening tabs)."""
        self.editor.blockSignals(True)
        self.editor.setPlainText(text)
        self.editor.blockSignals(False)
        self.is_modified = False
